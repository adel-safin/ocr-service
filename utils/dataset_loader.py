"""Утилита для загрузки и парсинга датасета"""
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import docx
import pandas as pd
import logging

logger = logging.getLogger(__name__)


class DatasetLoader:
    """Класс для загрузки пар документов из датасета"""
    
    def __init__(self, dataset_root: str):
        """
        Инициализация загрузчика датасета
        
        Args:
            dataset_root: Корневая директория датасета
        """
        self.dataset_root = Path(dataset_root)
        self.supported_image_exts = {'.pdf', '.jpg', '.jpeg', '.png', '.bmp'}
        self.supported_reference_exts = {'.doc', '.docx', '.txt', '.xlsx'}
    
    def find_document_pairs(self, document_type: Optional[str] = None) -> List[Dict[str, str]]:
        """
        Поиск пар документов (изображение + эталон)
        
        Args:
            document_type: Тип документа (например, "Акт АОСР") или None для всех
            
        Returns:
            Список словарей с путями к парам документов
        """
        pairs = []
        
        # Путь к директории с наборами документов
        documents_dir = self.dataset_root / "Наборы однотипных документов со сканами"
        
        if not documents_dir.exists():
            logger.warning(f"Директория не найдена: {documents_dir}")
            return pairs
        
        # Проход по типам документов
        for doc_type_dir in documents_dir.iterdir():
            if not doc_type_dir.is_dir():
                continue
            
            doc_type_name = doc_type_dir.name
            
            # Фильтрация по типу, если указан
            if document_type and document_type not in doc_type_name:
                continue
            
            # Поиск пар файлов в директории
            files = list(doc_type_dir.iterdir())
            
            # Группировка файлов по базовому имени
            file_groups = {}
            for file in files:
                base_name = self._get_base_name(file.stem)
                
                if base_name not in file_groups:
                    file_groups[base_name] = {'image': None, 'reference': None}
                
                ext = file.suffix.lower()
                if ext in self.supported_image_exts:
                    file_groups[base_name]['image'] = str(file)
                elif ext in self.supported_reference_exts:
                    file_groups[base_name]['reference'] = str(file)
            
            # Добавление пар, где есть и изображение, и эталон
            for base_name, files in file_groups.items():
                if files['image'] and files['reference']:
                    pairs.append({
                        'document_type': doc_type_name,
                        'base_name': base_name,
                        'image_path': files['image'],
                        'reference_path': files['reference'],
                        'image_ext': Path(files['image']).suffix.lower(),
                        'reference_ext': Path(files['reference']).suffix.lower()
                    })
        
        logger.info(f"Найдено {len(pairs)} пар документов")
        return pairs
    
    def _get_base_name(self, filename: str) -> str:
        """
        Извлечение базового имени файла (без номера)
        
        Args:
            filename: Имя файла
            
        Returns:
            Базовое имя
        """
        # Удаление номера в начале (например, "1 АОСР" -> "АОСР")
        parts = filename.split()
        if parts and parts[0].isdigit():
            return ' '.join(parts[1:])
        return filename
    
    def load_reference_text(self, reference_path: str) -> str:
        """
        Загрузка текста из эталонного файла
        
        Args:
            reference_path: Путь к эталонному файлу
            
        Returns:
            Текст из файла
        """
        ext = Path(reference_path).suffix.lower()
        
        try:
            if ext == '.txt':
                with open(reference_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif ext in ['.docx']:
                doc = docx.Document(reference_path)
                paragraphs = [p.text for p in doc.paragraphs]
                return '\n'.join(paragraphs)
            
            elif ext == '.doc':
                # Старый формат .doc требует специальной обработки
                # Пытаемся использовать python-docx, если не получается - пропускаем
                try:
                    doc = docx.Document(reference_path)
                    paragraphs = [p.text for p in doc.paragraphs]
                    return '\n'.join(paragraphs)
                except:
                    logger.warning(f"Не удалось прочитать .doc файл: {reference_path}")
                    return ""
            
            elif ext == '.xlsx':
                # Для Excel читаем все листы
                df = pd.read_excel(reference_path, sheet_name=None)
                texts = []
                for sheet_name, sheet_df in df.items():
                    texts.append(f"Лист: {sheet_name}")
                    texts.append(sheet_df.to_string())
                return '\n\n'.join(texts)
            
            else:
                logger.warning(f"Неподдерживаемый формат эталона: {ext}")
                return ""
        
        except Exception as e:
            logger.error(f"Ошибка при загрузке эталона {reference_path}: {str(e)}")
            return ""
    
    def get_all_document_types(self) -> List[str]:
        """
        Получение списка всех типов документов в датасете
        
        Returns:
            Список типов документов
        """
        documents_dir = self.dataset_root / "Наборы однотипных документов со сканами"
        
        if not documents_dir.exists():
            return []
        
        doc_types = [d.name for d in documents_dir.iterdir() if d.is_dir()]
        return sorted(doc_types)
    
    def create_training_pairs(self, document_type: Optional[str] = None) -> List[Dict]:
        """
        Создание пар для обучения (изображение, эталонный текст)
        
        Args:
            document_type: Тип документа или None для всех
            
        Returns:
            Список словарей с данными для обучения
        """
        pairs = self.find_document_pairs(document_type)
        training_data = []
        
        for pair in pairs:
            reference_text = self.load_reference_text(pair['reference_path'])
            
            if reference_text.strip():
                training_data.append({
                    'document_type': pair['document_type'],
                    'image_path': pair['image_path'],
                    'reference_text': reference_text,
                    'base_name': pair['base_name']
                })
        
        logger.info(f"Создано {len(training_data)} пар для обучения")
        return training_data
